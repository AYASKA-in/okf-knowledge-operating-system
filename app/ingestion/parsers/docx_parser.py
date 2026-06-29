import io
from pathlib import Path

from docx import Document

from app.ingestion.models import ParsedDocument, Section
from app.ingestion.parsers.base import DocumentParser


class DocxParser(DocumentParser):
    def parse(self, data: bytes, filename: str = "") -> ParsedDocument:
        title = Path(filename).stem if filename else "Untitled Document"
        doc = Document(io.BytesIO(data))

        sections: list[Section] = []
        current_title = title
        current_paragraphs: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                current_paragraphs.append("")
                continue
            if para.style.name.startswith("Heading"):
                if current_paragraphs:
                    sections.append(Section(
                        title=current_title,
                        text="\n".join(current_paragraphs).strip(),
                    ))
                current_title = text
                current_paragraphs = [text]
            else:
                current_paragraphs.append(text)

        if current_paragraphs:
            sections.append(Section(
                title=current_title,
                text="\n".join(current_paragraphs).strip(),
            ))

        return ParsedDocument(
            title=title,
            sections=[s for s in sections if s.text],
            source_type="docx",
            metadata={"filename": filename, "paragraphs": len(doc.paragraphs)},
        )
