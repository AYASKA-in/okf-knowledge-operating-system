"""Tests for document parsers."""
import io
import pytest

from app.ingestion.parsers import detect_format, parse_document


class TestDetectFormat:
    def test_pdf_by_magic(self):
        data = b"%PDF-1.4 some pdf content"
        assert detect_format(data) == "application/pdf"

    def test_pdf_by_extension(self):
        data = b"not actually pdf but has .pdf extension"
        assert detect_format(data, filename="doc.pdf") == "application/pdf"

    def test_html_by_magic(self):
        data = b"<html><head><title>Test</title></head></html>"
        assert detect_format(data) == "text/html"

    def test_html_doctype_by_magic(self):
        data = b"<!DOCTYPE html><html><head></head><body></body></html>"
        assert detect_format(data) == "text/html"

    def test_html_by_extension(self):
        data = b"garbage"
        assert detect_format(data, filename="page.htm") == "text/html"

    def test_docx_by_extension(self):
        data = b"PK\x03\x04" + b"\x00" * 26 + b"word/document.xml"
        assert detect_format(data, filename="report.docx") == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def test_xlsx_by_extension(self):
        data = b"PK\x03\x04" + b"\x00" * 26 + b"xl/workbook.xml"
        assert detect_format(data, filename="data.xlsx") == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

    def test_mime_hint_priority(self):
        data = b"some random data"
        assert detect_format(data, filename="doc.pdf", mime_hint="text/html") == "text/html"

    def test_unsupported_format_returns_empty(self):
        data = b"\x00\x01\x02\x03"
        assert detect_format(data) == ""

    def test_pdf_magic_in_html_body(self):
        data = b"<html><p>%PDF is not a real pdf</p></html>"
        assert detect_format(data) == "text/html"

    def test_empty_data(self):
        assert detect_format(b"") == ""


class TestPdfParser:
    def test_parse_simple_pdf(self):
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 700, "Hello World")
        c.drawString(100, 680, "This is a test PDF")
        c.save()
        data = buf.getvalue()

        doc = parse_document(data, filename="test.pdf")
        assert doc.title == "test"
        assert doc.source_type == "pdf"
        assert len(doc.sections) >= 1
        assert any("Hello" in s.text for s in doc.sections)
        assert "pages" in doc.metadata

    def test_parse_empty_pdf(self):
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.save()
        data = buf.getvalue()

        doc = parse_document(data, filename="empty.pdf")
        assert doc.title == "empty"
        assert doc.source_type == "pdf"


class TestDocxParser:
    def test_parse_simple_docx(self):
        from docx import Document
        buf = io.BytesIO()
        doc = Document()
        doc.add_heading("Introduction", 1)
        doc.add_paragraph("This is the intro paragraph.")
        doc.add_heading("Details", 2)
        doc.add_paragraph("More detailed content here.")
        doc.save(buf)
        data = buf.getvalue()

        doc_result = parse_document(data, filename="report.docx")
        assert doc_result.title == "report"
        assert doc_result.source_type == "docx"
        assert len(doc_result.sections) >= 1
        titles = [s.title for s in doc_result.sections]
        assert "Introduction" in titles or "report" in titles

    def test_parse_empty_docx(self):
        from docx import Document
        buf = io.BytesIO()
        Document().save(buf)
        data = buf.getvalue()

        doc_result = parse_document(data, filename="empty.docx")
        assert doc_result.source_type == "docx"


class TestXlsxParser:
    def test_parse_simple_xlsx(self):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws.append(["Name", "Age", "City"])
        ws.append(["Alice", "30", "New York"])
        ws.append(["Bob", "25", "London"])
        buf = io.BytesIO()
        wb.save(buf)
        data = buf.getvalue()

        doc = parse_document(data, filename="data.xlsx")
        assert doc.title == "data"
        assert doc.source_type == "xlsx"
        assert len(doc.sections) >= 1
        sheet_section = doc.sections[0]
        assert sheet_section.title == "Sheet1"
        assert "Alice" in sheet_section.text

    def test_parse_multi_sheet_xlsx(self):
        import openpyxl
        wb = openpyxl.Workbook()
        ws1 = wb.active
        ws1.title = "Summary"
        ws1.append(["Total", "100"])
        ws2 = wb.create_sheet("Details")
        ws2.append(["Item", "Count"])
        ws2.append(["X", "50"])
        ws2.append(["Y", "50"])
        buf = io.BytesIO()
        wb.save(buf)
        data = buf.getvalue()

        doc = parse_document(data, filename="report.xlsx")
        assert len(doc.sections) >= 2
        sheet_names = [s.title for s in doc.sections]
        assert "Summary" in sheet_names
        assert "Details" in sheet_names


class TestHtmlParser:
    def test_parse_simple_html(self):
        html = b"""<html><head><title>Test Page</title></head>
        <body><h1>Main Heading</h1><p>Some paragraph content.</p>
        <h2>Sub Section</h2><p>More details here.</p></body></html>"""
        doc = parse_document(html, filename="page.html")
        assert doc.title == "Test Page"
        assert doc.source_type == "html"
        assert len(doc.sections) >= 1

    def test_parse_html_without_title(self):
        html = b"<html><body><p>No title tag here</p></body></html>"
        doc = parse_document(html, filename="page.html")
        assert doc.title == "page"
        assert doc.source_type == "html"

    def test_html_strips_script_and_style(self):
        html = b"""<html><head><title>Clean</title></head>
        <body><h1>Heading</h1><script>alert('xss')</script>
        <style>.cls{color:red}</style><p>Real content</p></body></html>"""
        doc = parse_document(html, filename="page.html")
        full_text = " ".join(s.text for s in doc.sections)
        assert "alert" not in full_text
        assert ".cls" not in full_text
        assert "Real content" in full_text


class TestParseDocumentIntegration:
    def test_unsupported_format_raises(self):
        with pytest.raises(ValueError, match="Unsupported format"):
            parse_document(b"\x00\x01\x02\x03", filename="data.bin")

    def test_auto_detect_pdf_from_bytes(self):
        from reportlab.pdfgen import canvas
        buf = io.BytesIO()
        c = canvas.Canvas(buf)
        c.drawString(100, 700, "PDF content")
        c.save()
        doc = parse_document(buf.getvalue())
        assert doc.source_type == "pdf"

    def test_auto_detect_html_from_bytes(self):
        data = b"<html><body><p>Hello</p></body></html>"
        doc = parse_document(data)
        assert doc.source_type == "html"

    def test_docx_section_content(self):
        from docx import Document
        buf = io.BytesIO()
        doc = Document()
        doc.add_heading("Chapter 1", 1)
        doc.add_paragraph("Text content for chapter 1.")
        doc.save(buf)

        parsed = parse_document(buf.getvalue(), filename="book.docx")
        assert any("Chapter 1" in s.title for s in parsed.sections)

    def test_xlsx_pipe_separated_values(self):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["A", "B"])
        ws.append(["1", "2"])
        buf = io.BytesIO()
        wb.save(buf)

        parsed = parse_document(buf.getvalue(), filename="test.xlsx")
        assert " | " in parsed.sections[0].text
